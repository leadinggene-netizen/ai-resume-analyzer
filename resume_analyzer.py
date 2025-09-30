#!/usr/bin/env python3
"""
AI Agnets CV analyzer

"""
from logging import PlaceHolder
import re
import streamlit as st
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from datetime import datetime

API_BASE_URL="https://api.siliconflow.cn/v1"
TARGET_POSITION=["FrontEnd Developer","BackEnd Developer","FullStack Developer","Data Analytics Developer"]

def clean_text_for_pdf(text):
    """Clean text content to make it safe for PDF generation"""
    if not text:
        return ""
    
    # Remove HTML tags
    import re
    text = re.sub(r'<[^>]+>', '', text)
    
    # Replace markdown table syntax with plain text
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Skip table separator lines (|-------|-------|
        if re.match(r'^\s*\|[-\s\|]+\|\s*$', line):
            continue
        
        # Convert table rows to simple text
        if '|' in line and line.strip().startswith('|'):
            # Remove leading/trailing | and split by |
            cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
            if cells:  # Only process if there are cells
                cleaned_lines.append(' - '.join(filter(None, cells)))
        else:
            # Regular line - just clean up
            cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    
    # Replace common problematic characters
    text = text.replace('•', '* ')  # Replace bullet points
    text = text.replace('✅', '[YES] ')  # Replace checkmarks
    text = text.replace('❌', '[NO] ')   # Replace X marks
    text = text.replace('🔍', '[FOCUS] ')  # Replace magnifying glass
    text = text.replace('💡', '[TIP] ')   # Replace lightbulb
    text = text.replace('🌟', '[STAR] ')  # Replace star
    text = text.replace('🚀', '[GROWTH] ')  # Replace rocket
    
    # Remove any remaining emoji or special Unicode characters
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    
    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double
    text = re.sub(r' +', ' ', text)  # Multiple spaces to single
    
    return text.strip()

def extract_file_content(uploaded_file):
    if uploaded_file is None: 
        return None 
    
    try:
        if uploaded_file.type == "text/plain":
            # Handle TXT files
            file_content = uploaded_file.read()
            return file_content.decode("utf-8")
        
        elif uploaded_file.type == "application/pdf":
            # Handle PDF files
            pdf_reader = PdfReader(uploaded_file)
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            return text_content.strip()
        
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Handle DOCX files
            doc = Document(uploaded_file)
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            return text_content.strip()
        
        else:
            return f"Unsupported file type: {uploaded_file.type}. Please upload a TXT, PDF, or DOCX file."
    
    except Exception as e:
        return f"Error reading file: {str(e)}. Please ensure the file is not corrupted and try again."

def get_resume_content(uploaded_file,resume_text):
    if uploaded_file:
        return extract_file_content(uploaded_file)
    elif resume_text:
        return resume_text
    else:
        return None

def analyze_resume_with_ai(resume_text,target_postion,api_key):
    if not api_key or not api_key.strip():
        return "Enter your API key"

    client = OpenAI(api_key=api_key,base_url=API_BASE_URL)
    prompt = f"you are senior HR analyst assistant, taking account of the resume submitted by the candidate, you need to analyze the resume to validate if the candidate qualifies for the requirements listed in {target_postion}: \n based on{resume_text}please provide:1,overall assessment score (1-100);2,detailed analysis with suggestion for improvements; 3,list the candidate major outstandings, together with personalised advice on his/her future professional path"

    response = client.chat.completions.create(
        model="Qwen/Qwen3-Next-80B-A3B-Instruct",
        messages=[
            {"role":"system","content":"you act as an assistant like a senior professional to evaluate the candidate's resume with insightful analysis and advice"}
           ,{"role":"user","content":prompt}
        ],
        max_tokens = 1500
    )
    return response.choices[0].message.content

def analyze_job_posting(job_description, api_key):
    """Analyze job posting to extract key requirements and skills"""
    if not api_key or not api_key.strip():
        return "Enter your API key"

    client = OpenAI(api_key=api_key, base_url=API_BASE_URL)
    prompt = f"""You are a senior HR analyst. Analyze the following job posting and extract:
    1. Key required skills and technologies
    2. Preferred qualifications
    3. Soft skills and personality traits
    4. Experience level requirements
    5. Main responsibilities and duties
    6. Keywords that should appear in a tailored resume
    
    Job Posting:
    {job_description}
    
    Please provide a structured analysis that can be used to optimize a resume for this position."""

    response = client.chat.completions.create(
        model="Qwen/Qwen3-Next-80B-A3B-Instruct",
        messages=[
            {"role": "system", "content": "you act as an expert HR analyst who understands job requirements and can extract key information for resume optimization"},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500
    )
    return response.choices[0].message.content

def rewrite_resume_for_job(original_resume, job_analysis, evaluation_feedback, api_key):
    """Rewrite resume to better match job requirements based on evaluation feedback"""
    if not api_key or not api_key.strip():
        return "Enter your API key"

    client = OpenAI(api_key=api_key, base_url=API_BASE_URL)
    prompt = f"""You are a professional resume writer with expertise in ATS optimization and job matching.
    
    Based on the following information, rewrite the candidate's resume to better match the job requirements:
    
    ORIGINAL RESUME:
    {original_resume}
    
    JOB REQUIREMENTS ANALYSIS:
    {job_analysis}
    
    EVALUATION FEEDBACK:
    {evaluation_feedback}
    
    IMPORTANT: Provide ONLY the complete, polished resume content in standard resume format. Do NOT include:
    - Explanatory text or commentary
    - Analysis or evaluation remarks
    - Suggestions or recommendations
    - Meta-commentary about the changes made
    
    The output should be a clean, professional resume with:
    1. Contact information section
    2. Professional summary/objective
    3. Work experience with bullet points
    4. Education section
    5. Skills section
    6. Any relevant additional sections (certifications, projects, etc.)
    
    Optimize the resume by:
    - Using keywords from job requirements
    - Highlighting relevant experience and skills
    - Improving ATS compatibility
    - Quantifying achievements where possible
    - Tailoring content to the specific role
    
    Output the resume content ready for immediate use and submission."""

    response = client.chat.completions.create(
        model="Qwen/Qwen3-Next-80B-A3B-Instruct",
        messages=[
            {"role": "system", "content": "You are an expert resume writer who creates ATS-optimized, compelling resumes that match job requirements perfectly"},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2500
    )
    return response.choices[0].message.content

def generate_pdf_report(analysis_result, target_position, candidate_name="Candidate"):
    """Generate a PDF report from the analysis results"""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.darkblue,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.darkblue,
            spaceAfter=12,
            spaceBefore=20
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            textColor=colors.black
        )
        
        # Build the story
        story = []
        
        # Title
        story.append(Paragraph("AI Resume Analysis Report", title_style))
        story.append(Spacer(1, 20))
        
        # Header Information
        current_date = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(f"<b>Report Generated:</b> {current_date}", normal_style))
        story.append(Paragraph(f"<b>Target Position:</b> {target_position}", normal_style))
        story.append(Paragraph(f"<b>Candidate:</b> {candidate_name}", normal_style))
        story.append(Spacer(1, 30))
        
        # Analysis Results
        story.append(Paragraph("Analysis Results", heading_style))
        
        # Clean the analysis result text first
        cleaned_analysis = clean_text_for_pdf(analysis_result)
        
        # Split analysis result into paragraphs and format
        analysis_paragraphs = cleaned_analysis.split('\n\n')
        for paragraph in analysis_paragraphs:
            if paragraph.strip():
                # Process each paragraph safely
                paragraph_text = paragraph.strip()
                
                # Split long paragraphs at sentence boundaries to avoid overflow
                sentences = paragraph_text.split('. ')
                current_paragraph = ""
                
                for i, sentence in enumerate(sentences):
                    if len(current_paragraph + sentence) > 800:  # Prevent overly long paragraphs
                        if current_paragraph:
                            story.append(Paragraph(current_paragraph.strip() + '.', normal_style))
                            story.append(Spacer(1, 8))
                        current_paragraph = sentence
                    else:
                        if i < len(sentences) - 1:
                            current_paragraph += sentence + '. '
                        else:
                            current_paragraph += sentence
                
                # Add the final paragraph
                if current_paragraph.strip():
                    story.append(Paragraph(current_paragraph.strip(), normal_style))
                    story.append(Spacer(1, 12))
        
        # Footer
        story.append(Spacer(1, 50))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            alignment=1  # Center alignment
        )
        story.append(Paragraph("Generated by AI Resume Analyzer", footer_style))
        story.append(Paragraph("Contact: xzhu@cofomo.com", footer_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    except Exception as e:
        # If PDF generation fails, create a simple fallback PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        styles = getSampleStyleSheet()
        
        story = []
        story.append(Paragraph("AI Resume Analysis Report", styles['Title']))
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"Target Position: {target_position}", styles['Normal']))
        story.append(Spacer(1, 20))
        story.append(Paragraph("Analysis Results:", styles['Heading2']))
        story.append(Spacer(1, 10))
        
        # Add simplified analysis text
        simplified_text = clean_text_for_pdf(analysis_result)
        # Split into smaller chunks to avoid issues
        chunks = [simplified_text[i:i+1000] for i in range(0, len(simplified_text), 1000)]
        
        for chunk in chunks:
            if chunk.strip():
                story.append(Paragraph(chunk.strip(), styles['Normal']))
                story.append(Spacer(1, 10))
        
        doc.build(story)
        buffer.seek(0)
        return buffer

def generate_resume_docx(resume_content, filename="Resume"):
    """Generate a clean DOCX file from resume content"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create a new document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Clean the content - remove any remaining analysis commentary
        lines = resume_content.split('\n')
        clean_lines = []
        
        skip_phrases = [
            'analysis:', 'evaluation:', 'suggestion:', 'recommendation:', 
            'improvement:', 'based on', 'here is', 'here\'s', 'this resume',
            'the candidate', 'overall assessment', 'score:'
        ]
        
        for line in lines:
            line = line.strip()
            if line and not any(phrase.lower() in line.lower() for phrase in skip_phrases):
                clean_lines.append(line)
        
        # Process each line and add to document
        current_section = None
        
        for line in clean_lines:
            if not line:
                continue
                
            # Check if it's a section header (all caps or title case with colons)
            if (line.isupper() or 
                any(keyword in line.upper() for keyword in 
                    ['PROFESSIONAL SUMMARY', 'WORK EXPERIENCE', 'EDUCATION', 
                     'SKILLS', 'CONTACT', 'CERTIFICATIONS', 'PROJECTS', 'EXPERIENCE'])):
                # Add section header
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Inches(0.15)
                p.space_after = Inches(0.1)
                current_section = line
            elif line.startswith('-') or line.startswith('•') or line.startswith('*'):
                # Add bullet point
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                p.space_after = Inches(0.05)
            else:
                # Add regular paragraph
                p = doc.add_paragraph(line)
                p.space_after = Inches(0.08)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        # Fallback: return as plain text
        return io.BytesIO(resume_content.encode('utf-8'))

def generate_resume_txt(resume_content):
    """Generate a clean TXT file from resume content"""
    # Clean the content - remove any analysis commentary
    lines = resume_content.split('\n')
    clean_lines = []
    
    skip_phrases = [
        'analysis:', 'evaluation:', 'suggestion:', 'recommendation:', 
        'improvement:', 'based on', 'here is', 'here\'s', 'this resume',
        'the candidate', 'overall assessment', 'score:', 'commentary:',
        'note:', 'explanation:'
    ]
    
    for line in lines:
        line = line.strip()
        if line and not any(phrase.lower() in line.lower() for phrase in skip_phrases):
            clean_lines.append(line)
    
    clean_content = '\n'.join(clean_lines)
    return io.BytesIO(clean_content.encode('utf-8'))

def handle_analysis_click(uploaded_file,resume_text, target_position, api_key) :
    # Check usage limits for non-premium users
    if not st.session_state.is_premium_user:
        if st.session_state.usage_count >= 3:
            st.error("🚫 You've reached the free usage limit (3 analyses)!")
            st.markdown("### 🚀 Upgrade Options:")
            st.markdown("- **Premium Monthly ($9.99)**: Unlimited analyses + Advanced features")
            st.markdown("- **Pay Per Use ($2.99)**: Single analysis with premium features")
            st.markdown("[**📧 Contact us to upgrade**](mailto:upgrade@your-domain.com)")
            return None
    
    content=get_resume_content(uploaded_file,resume_text)
    if not content:
        st.warning("Please submit your resume file or input resume text")
        return
    
    # Increment usage counter for non-premium users
    if not st.session_state.is_premium_user:
        st.session_state.usage_count += 1
        
    with st.spinner("working on your file..."):
        analysis_result=analyze_resume_with_ai(content,target_position,api_key)
        st.session_state.analysis_result=analysis_result
        st.session_state.target_position=target_position
        st.session_state.original_resume_content=content  # Store original content
    st.success("Analysis Completed Successfully")
    return True

def main():
    st.set_page_config(page_title="AI Resum Analyzer",page_icon=":memo:",layout="wide")
    if "analysis_result" not in st.session_state:
        st.session_state.analysis_result=None
    if "job_analysis" not in st.session_state:
        st.session_state.job_analysis=None
    if "optimized_resume" not in st.session_state:
        st.session_state.optimized_resume=None
    if "original_resume_content" not in st.session_state:
        st.session_state.original_resume_content=None
    if "evaluation_in_progress" not in st.session_state:
        st.session_state.evaluation_in_progress=False
    if "usage_count" not in st.session_state:
        st.session_state.usage_count = 0
    if "is_premium_user" not in st.session_state:
        st.session_state.is_premium_user = False
    with st.sidebar:
        st.markdown("AI Resume Analyzer will help you to evaluate your candidateship based on your resume and your target position")
        
        # Usage tracking display
        if not st.session_state.is_premium_user:
            remaining_uses = max(0, 3 - st.session_state.usage_count)
            if remaining_uses > 0:
                st.info(f"🆓 Free analyses remaining: {remaining_uses}")
            else:
                st.error("🚫 Free limit reached! Upgrade for unlimited access.")
                if st.button("🚀 Upgrade to Premium - $9.99/month"):
                    st.markdown("[**👉 Get Premium Access**](mailto:upgrade@your-domain.com?subject=Premium%20Upgrade&body=I%20want%20to%20upgrade%20to%20premium%20access)")
        else:
            st.success("✨ Premium User - Unlimited Access")
        
        # Premium toggle for testing (remove in production)
        if st.checkbox("Enable Premium Features (Demo)"):
            st.session_state.is_premium_user = True
        
        st.markdown("---")
        
        api_key=st.text_input("Please enter your API key from SiliconFlow platform",type="password",placeholder="Please provide your API key",help="Get your API key from siliconflow.cn or use OpenAI API key")
        if api_key:
            st.success("API key entered!")
        else:
            st.warning("Please provide your API key.")
            st.markdown("**Don't have an API key?**")
            st.markdown("- [Get SiliconFlow API Key](https://siliconflow.cn) (Recommended)")
            st.markdown("- [Get OpenAI API Key](https://platform.openai.com)")
            st.markdown("- [📧 Contact us for managed service](mailto:support@your-domain.com)")
        st.markdown("---")
        st.markdown("### About Us")
        st.markdown("AI Resume Analyzer is empowered with OpenAI to help you to evaluate your candidateship based on your resume and your target position")
        st.markdown("### Contact Us")
        st.markdown("Contact us for any questions: xzhu@cofomo.com")
    st.title("AI Resume Analyzer and Optimizer :memo:")
    st.markdown("AI Resume Analyzer will help you to evaluate the candidate's qualifications pertaining to his /her target position")
   
    # Single column layout for better space utilization
    st.markdown("### Submit Your Resume File or Input Text Content")
    
    # Resume input section
    col_input1, col_input2 = st.columns([1, 1])
    
    with col_input1:
        target_position=st.selectbox("the position you are applying for",TARGET_POSITION,help="Select from list")
        
        st.markdown("#### Upload Resume File")
        st.info("📁 **Supported formats**: TXT, PDF, DOCX files")
        uploaded_file=st.file_uploader("Submit Resume File",type=["txt","pdf","docx"],help="Supports TXT, PDF, and DOCX formats")
        
        if uploaded_file is not None:
            st.success(f"✅ File uploaded: {uploaded_file.name} ({uploaded_file.type})")
    
    with col_input2:
        st.markdown("#### Or Enter Resume Text Manually")
        resume_text=st.text_area("Enter your text here",height=200, placeholder="Please enter your resume content")
    
    # Evaluate button (full width)
    if st.button("🔍 Evaluate Resume Now",type="primary",use_container_width=True):
        handle_analysis_click(uploaded_file,resume_text,target_position,api_key)
    
    # Results section (only show if there are results)
    if st.session_state.analysis_result:
        st.markdown("---")
        st.markdown("### 📊 Evaluation Results")
        st.markdown(f"**Target Position:** {st.session_state.get('target_position','Not Selected')}")
        
        # Analysis results in expandable container for better readability
        with st.expander("📋 Detailed Analysis Report", expanded=True):
            st.markdown(st.session_state.analysis_result)

        # Action buttons
        st.markdown("#### 📄 Download Analysis Report")
        col_a, col_b = st.columns(2)
        with col_a:
                if st.button("📄 Download PDF Report",use_container_width=True):
                    try:
                        # Generate PDF report
                        with st.spinner("Generating PDF report..."):
                            pdf_buffer = generate_pdf_report(
                                st.session_state.analysis_result, 
                                st.session_state.get('target_position', 'Not Selected')
                            )
                        
                        # Create download button
                        current_date = datetime.now().strftime("%Y%m%d_%H%M%S")
                        filename = f"Resume_Analysis_Report_{current_date}.pdf"
                        
                        st.download_button(
                            label="📥 Click to Download PDF",
                            data=pdf_buffer.getvalue(),
                            file_name=filename,
                            mime="application/pdf",
                            use_container_width=True
                        )
                        st.success("✅ PDF report generated successfully!")
                    except Exception as e:
                        st.error(f"❌ Error generating PDF: {str(e)}")
                        # Show debug info in expander
                        with st.expander("Debug Information"):
                            st.text(f"Error details: {str(e)}")
                            st.text("Analysis result preview:")
                            st.text(st.session_state.analysis_result[:500] + "..." if len(st.session_state.analysis_result) > 500 else st.session_state.analysis_result)
        
        with col_b:
            if st.button("🔄 Evaluate Again",use_container_width=True):
                st.session_state.analysis_result=None
                st.session_state.job_analysis=None
                st.session_state.optimized_resume=None
                st.rerun()
        
        # Job Optimization Section
        st.markdown("---")
        st.markdown("### 🎯 Resume Optimization for Specific Job")
        
        # Premium feature check
        if not st.session_state.is_premium_user:
            st.warning("⭐ **Premium Feature**: Job optimization requires premium access")
            col_upgrade1, col_upgrade2 = st.columns([1, 1])
            with col_upgrade1:
                if st.button("🚀 Upgrade to Premium ($9.99/month)", use_container_width=True):
                    st.markdown("[**Get Premium Access**](mailto:upgrade@your-domain.com?subject=Premium%20Upgrade)")
            with col_upgrade2:
                if st.button("💳 Pay Per Use ($2.99)", use_container_width=True):
                    st.markdown("[**One-time Payment**](mailto:support@your-domain.com?subject=Pay%20Per%20Use)")
        else:
            st.info("💡 Upload or paste a job posting to get a tailored, optimized version of your resume!")
        
        # Job posting input (only show for premium users)
        if st.session_state.is_premium_user:
            col_job1, col_job2 = st.columns([1, 1])
            
            with col_job1:
                st.markdown("#### Upload Job Posting File")
                job_file = st.file_uploader("Upload Job Posting", type=["txt", "pdf", "docx"], 
                                          help="Upload the job posting document", key="job_file")
                if job_file is not None:
                    st.success(f"✅ Job posting uploaded: {job_file.name}")
            
            with col_job2:
                st.markdown("#### Or Paste Job Description")
                job_description = st.text_area("Paste job description here", height=200, 
                                              placeholder="Paste the complete job posting including requirements, responsibilities, and qualifications...",
                                              key="job_desc")
        else:
            # Show locked state for non-premium users
            st.text_input("🔒 Upload Job Posting File (Premium Only)", disabled=True, placeholder="Upgrade to premium to unlock this feature")
            st.text_area("🔒 Or Paste Job Description (Premium Only)", disabled=True, height=200, 
                        placeholder="Upgrade to premium to access job optimization features...")
        
        # Analyze job and optimize resume button (premium only)
        if st.session_state.is_premium_user:
            optimize_button = st.button("🚀 Optimize Resume for This Job", type="primary", use_container_width=True)
        else:
            st.button("🔒 Optimize Resume for This Job (Premium Only)", disabled=True, use_container_width=True)
            optimize_button = False
            
        if optimize_button:
            # Get job content (initialize variables for premium users)
            if st.session_state.is_premium_user:
                if 'job_file' in locals() and job_file:
                    job_content = extract_file_content(job_file)
                elif 'job_description' in locals() and job_description:
                    job_content = job_description
                else:
                    st.error("Please upload a job posting file or paste the job description")
                    job_content = None
            
            if job_content and st.session_state.original_resume_content:
                with st.spinner("Analyzing job requirements..."):
                    # Analyze job posting
                    job_analysis = analyze_job_posting(job_content, api_key)
                    st.session_state.job_analysis = job_analysis
                
                with st.spinner("Optimizing your resume for this job..."):
                    # Rewrite resume based on job and evaluation
                    optimized_resume = rewrite_resume_for_job(
                        st.session_state.original_resume_content,
                        job_analysis,
                        st.session_state.analysis_result,
                        api_key
                    )
                    st.session_state.optimized_resume = optimized_resume
                
                st.success("🎉 Resume optimized successfully!")
        
        # Display optimization results
        if st.session_state.job_analysis and st.session_state.optimized_resume:
                st.markdown("---")
                st.markdown("### 📋 Job Analysis & Optimized Resume")
                
                # Job analysis in expandable section
                with st.expander("🔍 Job Requirements Analysis", expanded=False):
                    st.markdown(st.session_state.job_analysis)
                
                # Optimized resume
                st.markdown("#### 📄 Your Optimized Resume")
                st.markdown(st.session_state.optimized_resume)
                
                # Download optimized resume in editable formats
                st.info("💡 Download your optimized resume in editable format for further customization")
                col_opt1, col_opt2, col_opt3 = st.columns(3)
                
                with col_opt1:
                    if st.button("� Download DOCX", use_container_width=True, help="Editable Word document"):
                        try:
                            with st.spinner("Generating DOCX resume..."):
                                docx_buffer = generate_resume_docx(st.session_state.optimized_resume)
                            
                            current_date = datetime.now().strftime("%Y%m%d_%H%M%S")
                            filename = f"Optimized_Resume_{current_date}.docx"
                            
                            st.download_button(
                                label="📥 Download DOCX",
                                data=docx_buffer.getvalue(),
                                file_name=filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="optimized_resume_docx"
                            )
                            st.success("✅ DOCX resume ready!")
                        except Exception as e:
                            st.error(f"❌ Error generating DOCX: {str(e)}")
                
                with col_opt2:
                    if st.button("📝 Download TXT", use_container_width=True, help="Plain text format"):
                        try:
                            txt_buffer = generate_resume_txt(st.session_state.optimized_resume)
                            
                            current_date = datetime.now().strftime("%Y%m%d_%H%M%S")
                            filename = f"Optimized_Resume_{current_date}.txt"
                            
                            st.download_button(
                                label="📥 Download TXT",
                                data=txt_buffer.getvalue(),
                                file_name=filename,
                                mime="text/plain",
                                use_container_width=True,
                                key="optimized_resume_txt"
                            )
                            st.success("✅ TXT resume ready!")
                        except Exception as e:
                            st.error(f"❌ Error generating TXT: {str(e)}")
                
                with col_opt3:
                    if st.button("🔄 Try Different Job", use_container_width=True):
                        st.session_state.job_analysis = None
                        st.session_state.optimized_resume = None
                        st.rerun()
    else:
        # Show information when no evaluation has been completed yet
        st.info("""
        ### 🎯 AI Resume Analyzer Features:
        
        **Step 1:** Upload your resume (TXT, PDF, or DOCX) or paste the text content
        
        **Step 2:** Select your target position and click "Evaluate Resume Now"
        
        **Step 3:** Review your detailed analysis with personalized feedback
        
        **Step 4:** (Optional) Upload a specific job posting to get an optimized resume
        
        **Benefits:**
        - Personalized analysis based on your target position
        - AI-powered scoring and detailed feedback
        - Professional career path recommendations
        - Job-specific resume optimization
        - Multiple download formats (PDF report, DOCX/TXT resume)
        """)

if __name__=="__main__":
    main()
    




    